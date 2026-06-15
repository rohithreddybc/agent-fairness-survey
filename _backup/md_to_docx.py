#!/usr/bin/env python3
"""Markdown -> .docx for the survey (python-docx; no pandoc).
Handles headings, markdown tables, bold/italic/code, bullet lists, and renders
\\cite{a,b} as [a, b]. Appends a References section built from references.bib for
the cited keys. Usage: python md_to_docx.py AGENT_FAIRNESS_SURVEY.md out.docx"""
from __future__ import annotations
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "AGENT_FAIRNESS_SURVEY.md"
OUT = Path(sys.argv[2]) if len(sys.argv) > 2 else HERE / "AGENT_FAIRNESS_SURVEY.docx"

md = SRC.read_text(encoding="utf-8")
cited = []  # preserve order of first appearance

def cites_to_text(s: str) -> str:
    def repl(m):
        keys = [k.strip() for k in m.group(1).split(",")]
        for k in keys:
            if k not in cited:
                cited.append(k)
        return "[" + ", ".join(keys) + "]"
    return re.sub(r"\\cite[tp]?\{([^}]+)\}", repl, s)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

def add_runs(par, text):
    text = cites_to_text(text)
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            r = par.add_run(part[1:-1]); r.italic = True
        else:
            par.add_run(part)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

lines = md.splitlines()
i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    # table block
    if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|?$", lines[i+1].strip()):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        header = rows[0]; body = [r for r in rows[2:]]
        t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
        for j, c in enumerate(header):
            cell = t.rows[0].cells[j]; cell.paragraphs[0].clear()
            add_runs(cell.paragraphs[0], c)
            for r in cell.paragraphs[0].runs: r.bold = True
        for br in body:
            cells = t.add_row().cells
            for j, c in enumerate(br[:len(header)]):
                cells[j].paragraphs[0].clear(); add_runs(cells[j].paragraphs[0], c)
        doc.add_paragraph()
        continue
    if not s:
        i += 1; continue
    if s == "---" or s.startswith("<!--"):
        i += 1; continue
    m = re.match(r"^(#{1,6})\s+(.*)$", s)
    if m:
        lvl = len(m.group(1)); txt = cites_to_text(m.group(2)).strip("*")
        if lvl == 1:
            h = doc.add_heading("", level=0); add_runs(h, m.group(2))
        else:
            h = doc.add_heading("", level=min(lvl-1, 4)); h.clear(); add_runs(h, m.group(2))
        i += 1; continue
    if re.match(r"^[-*]\s+", s):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r"^[-*]\s+", "", s))
        i += 1; continue
    if s.startswith(">"):
        p = doc.add_paragraph(); p.style = "Intense Quote"; add_runs(p, s.lstrip("> ").strip())
        i += 1; continue
    p = doc.add_paragraph(); add_runs(p, s)
    i += 1

# References from references.bib for cited keys
bib = (HERE / "references.bib").read_text(encoding="utf-8")
entries = {}
for m in re.finditer(r"@\w+\{([^,]+),(.*?)\n\}", bib, re.S):
    key = m.group(1).strip(); body = m.group(2)
    def f(name):
        mm = re.search(rf"{name}\s*=\s*\{{(.*?)\}}", body, re.S)
        return re.sub(r"\s+", " ", mm.group(1)).strip() if mm else ""
    entries[key] = (f("author"), f("year"), f("title"), f("journal") or f("booktitle"),
                    f("eprint"), f("doi"))

doc.add_page_break()
doc.add_heading("References", level=1)
for k in cited:
    if k not in entries:
        continue
    au, yr, ti, ve, ep, doi = entries[k]
    bits = [b for b in [au, f"({yr})" if yr else "", ti + ("." if ti and not ti.endswith('.') else ""),
                        ve, ("arXiv:" + ep) if ep else "", ("doi:" + doi) if doi else ""] if b]
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18); p.add_run(". ".join(bits))

doc.save(OUT)
print(f"wrote {OUT.name}: {len(cited)} cited refs, {len(doc.paragraphs)} paragraphs")
